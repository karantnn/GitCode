"""
JSON to Word Document Converter for Trading Agents Analysis

Converts JSON analysis files into professionally formatted Word documents.
Supports single or batch conversion with customizable styling.
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Union
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, fill_color):
    """Set cell background color in table"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_heading_with_line(doc, text, level=1):
    """Add heading with line separator"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if level == 1:
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.space_before = Pt(6)
        p_format.space_after = Pt(6)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '24')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '4472C4')
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    return heading


def add_metadata_table(doc, metadata: Dict):
    """Add metadata as a formatted table"""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Property'
    header_cells[1].text = 'Value'
    
    # Style header
    for cell in header_cells:
        set_cell_background(cell, 'E7E6E6')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add metadata rows
    for key, value in metadata.items():
        if key != 'analysis':  # Skip analysis content in metadata table
            row_cells = table.add_row().cells
            row_cells[0].text = key.replace('_', ' ').title()
            row_cells[1].text = str(value)[:100]  # Truncate long values
    
    doc.add_paragraph()


def format_analysis_text(text: str) -> str:
    """Format analysis text for better readability"""
    if not text:
        return ""
    
    # Clean up markdown-style headers
    lines = text.split('\n')
    formatted = []
    
    for line in lines:
        # Keep lines as-is (they may contain markdown)
        formatted.append(line)
    
    return '\n'.join(formatted)


def add_analysis_section(doc, analysis_text: str):
    """Add formatted analysis section"""
    if not analysis_text:
        return
    
    doc.add_heading('Analysis Details', level=2)
    
    # Parse and format analysis
    lines = analysis_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Handle markdown headers
        if line.startswith('###'):
            text = line.replace('###', '').strip()
            p = doc.add_heading(text, level=3)
            p_format = p.paragraph_format
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(3)
        elif line.startswith('##'):
            text = line.replace('##', '').strip()
            p = doc.add_heading(text, level=2)
            p_format = p.paragraph_format
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(3)
        elif line.startswith('#'):
            text = line.replace('#', '').strip()
            p = doc.add_heading(text, level=1)
            p_format = p.paragraph_format
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(3)
        elif line.startswith('-') or line.startswith('•'):
            # Bullet point
            text = line.lstrip('-•').strip()
            p = doc.add_paragraph(text, style='List Bullet')
        elif line.startswith('*'):
            # Alternative bullet
            text = line.lstrip('*').strip()
            p = doc.add_paragraph(text, style='List Bullet')
        else:
            # Regular paragraph
            p = doc.add_paragraph(line)
            p_format = p.paragraph_format
            p_format.space_after = Pt(6)
            p_format.line_spacing = 1.15


def json_to_word(json_file: Union[str, Path], output_file: Union[str, Path] = None) -> Path:
    """
    Convert a single JSON file to a Word document
    
    Args:
        json_file: Path to JSON file
        output_file: Path for output Word document (auto-generated if not provided)
    
    Returns:
        Path to created Word document
    """
    json_path = Path(json_file)
    
    if not json_path.exists():
        raise FileNotFoundError(f"JSON file not found: {json_file}")
    
    # Load JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Determine output path
    if output_file is None:
        output_file = json_path.parent / f"{json_path.stem}.docx"
    else:
        output_file = Path(output_file)
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Document title
    title = doc.add_heading('Trading Agents Analysis Report', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with stock info
    agent_name = data.get('agent_name', 'Unknown Agent')
    ticker = data.get('ticker', 'N/A')
    date = data.get('date', 'N/A')
    
    subtitle = doc.add_paragraph(f"{agent_name} - {ticker} ({date})")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(12)
    subtitle_format.font.italic = True
    subtitle_format.font.color.rgb = RGBColor(89, 89, 89)
    
    doc.add_paragraph()  # Spacing
    
    # Metadata section
    add_heading_with_line(doc, 'Metadata', level=1)
    metadata = {k: v for k, v in data.items() if k != 'analysis'}
    add_metadata_table(doc, metadata)
    
    # Analysis section
    if 'analysis' in data and data['analysis']:
        add_heading_with_line(doc, 'Analysis Report', level=1)
        add_analysis_section(doc, data['analysis'])
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph(
        f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
        f"Source: {json_path.name}"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_format = footer_para.runs[0]
    footer_format.font.size = Pt(9)
    footer_format.font.italic = True
    footer_format.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    doc.save(output_file)
    
    return output_file


def batch_json_to_word(
    input_dir: Union[str, Path],
    output_dir: Union[str, Path] = None,
    pattern: str = "*.json",
    combine: bool = False
) -> List[Path]:
    """
    Convert multiple JSON files to Word documents
    
    Args:
        input_dir: Directory containing JSON files
        output_dir: Directory for output Word documents (defaults to input_dir)
        pattern: File pattern to match (default: "*.json")
        combine: If True, create single document with all analyses
    
    Returns:
        List of created Word document paths
    """
    input_dir = Path(input_dir)
    output_dir = Path(output_dir) if output_dir else input_dir
    
    if not input_dir.exists():
        raise FileNotFoundError(f"Input directory not found: {input_dir}")
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Find all JSON files
    json_files = sorted(input_dir.glob(pattern))
    
    if not json_files:
        print(f"No JSON files matching '{pattern}' found in {input_dir}")
        return []
    
    created_files = []
    
    if combine:
        # Create single combined document
        doc = Document()
        
        # Document title
        title = doc.add_heading('Trading Agents Analysis Report - Combined', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        for json_file in json_files:
            try:
                with open(json_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Add section for each analysis
                agent_name = data.get('agent_name', 'Unknown')
                ticker = data.get('ticker', 'N/A')
                date = data.get('date', 'N/A')
                
                add_heading_with_line(doc, f"{agent_name} - {ticker} ({date})", level=1)
                
                # Metadata table
                metadata = {k: v for k, v in data.items() if k != 'analysis'}
                add_metadata_table(doc, metadata)
                
                # Analysis
                if 'analysis' in data and data['analysis']:
                    add_heading_with_line(doc, 'Analysis Details', level=2)
                    add_analysis_section(doc, data['analysis'])
                
                # Page break between analyses
                doc.add_page_break()
                
                print(f"✓ Added {agent_name} to combined document")
            
            except Exception as e:
                print(f"✗ Error processing {json_file.name}: {e}")
        
        # Save combined document
        output_file = output_dir / "Combined_Analysis.docx"
        doc.save(output_file)
        created_files.append(output_file)
        print(f"\n✓ Combined document created: {output_file}")
    
    else:
        # Create individual documents
        for json_file in json_files:
            try:
                output_file = output_dir / f"{json_file.stem}.docx"
                result = json_to_word(json_file, output_file)
                created_files.append(result)
                print(f"✓ Created: {result.name}")
            
            except Exception as e:
                print(f"✗ Error processing {json_file.name}: {e}")
    
    return created_files


def main():
    """Command-line interface"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Convert JSON analysis files to formatted Word documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert single JSON to Word
  python json_to_word.py results/INTC/2025-12-25/fundamentals_analysis.json
  
  # Convert with custom output name
  python json_to_word.py results/INTC/2025-12-25/analysis.json output.docx
  
  # Batch convert all JSON files in directory
  python json_to_word.py results/INTC/2025-12-25 --batch
  
  # Batch convert and combine into single document
  python json_to_word.py results/INTC/2025-12-25 --batch --combine
  
  # Batch convert to different output directory
  python json_to_word.py results/INTC/2025-12-25 --batch --output-dir reports
        """
    )
    
    parser.add_argument(
        'input',
        help='Input JSON file or directory'
    )
    
    parser.add_argument(
        '-o', '--output',
        dest='output',
        help='Output file path (for single file) or directory (for batch)'
    )
    
    parser.add_argument(
        '-b', '--batch',
        action='store_true',
        help='Process all JSON files in directory'
    )
    
    parser.add_argument(
        '-c', '--combine',
        action='store_true',
        help='Combine all analyses into single document (with --batch)'
    )
    
    parser.add_argument(
        '-p', '--pattern',
        default='*.json',
        help='File pattern for batch mode (default: *.json)'
    )
    
    args = parser.parse_args()
    
    try:
        if args.batch:
            # Batch mode
            print(f"Processing directory: {args.input}")
            results = batch_json_to_word(
                args.input,
                output_dir=args.output,
                pattern=args.pattern,
                combine=args.combine
            )
            print(f"\nCreated {len(results)} document(s)")
        else:
            # Single file mode
            print(f"Converting: {args.input}")
            result = json_to_word(args.input, args.output)
            print(f"✓ Created: {result}")
    
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
